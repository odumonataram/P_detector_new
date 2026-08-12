"""
Sample Test Script for Plagiarism Detection System
This script demonstrates basic functionality and creates test documents
"""

import os
from docx import Document
from datetime import datetime


def create_sample_documents():
    """Create sample DOCX documents for testing"""
    
    # Sample Document 1 - Original
    doc1 = Document()
    doc1.add_heading('The Impact of Climate Change on Biodiversity', 0)
    doc1.add_paragraph(
        'Climate change represents one of the most significant threats to global biodiversity. '
        'Rising temperatures and changing precipitation patterns are altering ecosystems worldwide. '
        'Species are being forced to adapt, migrate, or face extinction. The rapid pace of change '
        'exceeds the adaptive capacity of many organisms, leading to widespread ecological disruption.'
    )
    doc1.add_paragraph(
        'Marine ecosystems are particularly vulnerable to climate change. Ocean acidification, '
        'caused by increased carbon dioxide absorption, threatens coral reefs and shellfish populations. '
        'Rising sea temperatures lead to coral bleaching events, devastating these critical habitats. '
        'The loss of coral reefs has cascading effects on countless species that depend on them.'
    )
    doc1.add_paragraph(
        'Terrestrial ecosystems face similar challenges. Changes in temperature and rainfall patterns '
        'disrupt plant phenology and animal migration patterns. Species with limited dispersal abilities '
        'or specific habitat requirements are at greatest risk. Conservation efforts must adapt to these '
        'changing conditions to remain effective.'
    )
    doc1.add_heading('References', 1)
    doc1.add_paragraph('Smith, J. (2023). Climate Change and Biodiversity. Nature Press.')
    
    doc1.save('sample_doc1_original.docx')
    print("✓ Created sample_doc1_original.docx")
    
    # Sample Document 2 - Partially Plagiarized
    doc2 = Document()
    doc2.add_heading('Effects of Global Warming on Wildlife', 0)
    doc2.add_paragraph(
        'Global warming has emerged as a critical environmental challenge affecting wildlife worldwide. '
        'Rising temperatures and changing precipitation patterns are altering ecosystems worldwide. '
        'Many species are struggling to adapt to these rapid environmental changes.'
    )
    doc2.add_paragraph(
        'Ocean ecosystems are experiencing severe impacts from climate change. Ocean acidification, '
        'caused by increased carbon dioxide absorption, threatens coral reefs and shellfish populations. '
        'These changes are fundamentally altering marine food webs and ecosystem dynamics.'
    )
    doc2.add_paragraph(
        'On land, wildlife populations are responding to shifting climate patterns. Migration routes are '
        'changing, breeding seasons are shifting, and some species are moving to higher elevations. '
        'Conservationists are working to develop strategies that account for these ongoing changes.'
    )
    doc2.add_heading('Bibliography', 1)
    doc2.add_paragraph('Jones, A. (2024). Wildlife in a Warming World. Science Publishers.')
    
    doc2.save('sample_doc2_partial.docx')
    print("✓ Created sample_doc2_partial.docx")
    
    # Sample Document 3 - Mostly Original
    doc3 = Document()
    doc3.add_heading('Renewable Energy Solutions for Sustainable Development', 0)
    doc3.add_paragraph(
        'The transition to renewable energy sources is essential for sustainable development. '
        'Solar, wind, and hydroelectric power offer clean alternatives to fossil fuels. '
        'These technologies have matured significantly, becoming economically competitive with '
        'traditional energy sources in many regions.'
    )
    doc3.add_paragraph(
        'Solar energy technology has advanced rapidly in recent years. Photovoltaic cell efficiency '
        'has improved while costs have decreased dramatically. Solar installations are now viable '
        'in diverse climates and settings, from large-scale solar farms to residential rooftops.'
    )
    doc3.add_paragraph(
        'Wind energy represents another crucial component of the renewable energy mix. Modern wind '
        'turbines generate significant power with minimal environmental impact. Offshore wind farms '
        'are expanding globally, tapping into consistent ocean winds to generate clean electricity.'
    )
    doc3.add_heading('References', 1)
    doc3.add_paragraph('Brown, K. (2024). Renewable Energy Technologies. Green Press.')
    
    doc3.save('sample_doc3_original.docx')
    print("✓ Created sample_doc3_original.docx")
    
    print("\n📄 Sample documents created successfully!")
    print("\nNext steps:")
    print("1. Run the Flask application: python app.py")
    print("2. Open browser to: http://localhost:5000")
    print("3. Upload these sample documents to test the system")
    print("4. Try uploading sample_doc2_partial.docx after sample_doc1_original.docx")
    print("   to see plagiarism detection in action!")


def create_test_directory():
    """Create a test directory for sample documents"""
    test_dir = 'sample_documents'
    if not os.path.exists(test_dir):
        os.makedirs(test_dir)
    os.chdir(test_dir)


if __name__ == '__main__':
    print("=" * 60)
    print("Plagiarism Detection System - Sample Document Generator")
    print("=" * 60)
    print()
    
    create_test_directory()
    create_sample_documents()
    
    print("\n" + "=" * 60)
    print("Sample documents are ready for testing!")
    print("=" * 60)
